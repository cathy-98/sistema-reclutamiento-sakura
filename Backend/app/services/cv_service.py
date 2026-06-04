import io
import re
from datetime import datetime, date
from pypdf import PdfReader
from fastapi import HTTPException
from app.schemas import CandidatoCreate

def extraer_texto_pdf(contenido_archivo: bytes) -> str:
    """
    Lee los bytes de un archivo PDF en memoria, extrae todo su texto plano utilizando pypdf
    y limpia de forma agresiva los caracteres NUL (0x00) que corrompen PostgreSQL.
    """
    try:
        pdf_file = io.BytesIO(contenido_archivo)
        reader = PdfReader(pdf_file)
        texto_completo = []
        
        for pagina in reader.pages:
            texto_pagina = pagina.extract_text()
            if texto_pagina:
                # CAMBIO CLÍTICO: Limpiamos los caracteres nulos invisibles (\x00)
                texto_limpio = texto_pagina.replace('\x00', '')
                texto_completo.append(texto_limpio)
                
        return "\n".join(texto_completo)
    except Exception as e:
        raise HTTPException(
            status_code=400,
            detail=f"No se pudo extraer el texto del archivo PDF: {str(e)}"
        )

def extraer_texto_docx(contenido_archivo: bytes) -> str:
    """
    Lee los bytes de un archivo DOCX en memoria y extrae todo su texto plano utilizando python-docx.
    """
    try:
        import docx
        docx_file = io.BytesIO(contenido_archivo)
        doc = docx.Document(docx_file)
        texto_completo = []
        
        for parrafo in doc.paragraphs:
            if parrafo.text:
                texto_completo.append(parrafo.text)
                
        for tabla in doc.tables:
            for fila in tabla.rows:
                for celda in fila.cells:
                    if celda.text:
                        texto_completo.append(celda.text)
                        
        return "\n".join(texto_completo)
    except ImportError:
        raise HTTPException(
            status_code=502,
            detail="La librería 'python-docx' no está instalada. Por favor ejecute: pip install python-docx"
        )
    except Exception as e:
        raise HTTPException(
            status_code=400,
            detail=f"No se pudo extraer el texto del archivo DOCX: {str(e)}"
        )

async def procesar_cv_con_python(texto_cv: str) -> CandidatoCreate:
    """
    Procesa el texto de forma 100% local usando expresiones regulares y lógica avanzada de Python.
    Repara el efecto Canva de forma global y utiliza una extracción de nombres altamente selectiva.
    """
    # =========================================================================
    # 0. REPARACIÓN GLOBAL DEL EFECTO CANVA / FIGMA
    # =========================================================================
    # Si detectamos que hay muchas letras sueltas separadas por espacios en el documento,
    # reparamos el texto completo antes de pasarlo a los extractores.
    texto_normalizado = texto_cv
    if len(re.findall(r'\b\w \w\b', texto_cv)) > 10:
        lineas_reparadas = []
        for linea in texto_cv.split('\n'):
            # Une letras solas consecutivas: "C a t h e r i n e" -> "Catherine"
            # Pero respeta espacios entre palabras reales
            linea_reparada = re.sub(r'(?<=\b\w)\s+(?=\w\b)', '', linea)
            # Une números separados: "5 6 9" -> "569"
            linea_reparada = re.sub(r'(?<=\d)\s+(?=\d)', '', linea_reparada)
            lineas_reparadas.append(linea_reparada)
        texto_normalizado = "\n".join(lineas_reparadas)

    # =========================================================================
    # 1. EXTRAER CORREO ELECTRÓNICO
    # =========================================================================
    pattern_email = r'[a-zA-Z0-9-_\.]+@[a-zA-Z0-9-_\.]+\.[a-zA-Z]{2,5}'
    match_email = re.search(pattern_email, texto_normalizado)
    if match_email:
        correo = match_email.group(0).strip().lower()
    else:
        correo = None
        for linea in texto_cv.split('\n'):
            if "@" in linea:
                linea_sin_espacios = linea.replace(" ", "")
                match_linea = re.search(pattern_email, linea_sin_espacios)
                if match_linea:
                    correo = match_linea.group(0).strip().lower()
                    break

    if not correo:
        raise HTTPException(
            status_code=422,
            detail="No se pudo procesar el CV: No se encontró un correo electrónico válido."
        )

    # =========================================================================
    # 2. EXTRAER TELÉFONO DE CONTACTO
    # =========================================================================
    telefono = None
    match_etiqueta = re.search(r'(?:teléfono|telefono|celular|fono)\s*:\s*(\+?[\d\s-]{7,15})', texto_normalizado, re.IGNORECASE)
    if match_etiqueta:
        limpio = re.sub(r'[\s\n\r-]', '', match_etiqueta.group(1))
        if len(limpio) >= 8:
            telefono = limpio
            
    if not telefono:
        pattern_tel_flexible = r'(\+?\s*56\s*9?[\d\s-]{7,15}|\b9\s*[\d\s-]{7,12}\b)'
        match_tel = re.search(pattern_tel_flexible, texto_normalizado)
        if match_tel:
            telefono = re.sub(r'[\s\n\r-]', '', match_tel.group(1))

    if telefono:
        if telefono.startswith("56"):
            telefono = "+" + telefono
        elif telefono.startswith("9") and len(telefono) == 9:
            telefono = "+56" + telefono
        elif not telefono.startswith("+") and telefono.startswith("56"):
            telefono = "+" + telefono

    # =========================================================================
    # 3. EXTRAER RUT
    # =========================================================================
    rut = None
    pattern_rut = r'\b(\d{1,2}(?:\.?\d{3}){2}-[\dkK])\b'
    match_rut = re.search(pattern_rut, texto_normalizado)
    if match_rut:
        rut = match_rut.group(1).strip().upper()

    # =========================================================================
    # 4. EXTRAER FECHA DE NACIMIENTO
    # =========================================================================
    fecha_obj: date | None = None
    fecha_texto = None
    lineas_fecha = [l for l in texto_normalizado.split('\n') if any(k in l.lower() for k in ["nacimiento", "f.n", "fecha de nac", "nacido"])]
    pattern_fecha_num = r'(\b\d{1,2}[-/\.]\d{1,2}[-/\.]\d{4}\b)'
    pattern_fecha_texto = r'(\b\d{1,2}\s+(?:de\s+)?(?:enero|febrero|marzo|abril|mayo|junio|julio|agosto|septiembre|octubre|noviembre|diciembre)\s+(?:de\s+)?\d{4}\b)'
    
    for linea in lineas_fecha:
        match_f = re.search(pattern_fecha_num, linea) or re.search(pattern_fecha_texto, linea, re.IGNORECASE)
        if match_f:
            fecha_texto = match_f.group(1).strip()
            break
            
    if not fecha_texto:
        match_f_global = re.search(pattern_fecha_num, texto_normalizado) or re.search(pattern_fecha_texto, texto_normalizado, re.IGNORECASE)
        if match_f_global:
            fecha_texto = match_f_global.group(1).strip()

    if fecha_texto:
        fecha_normalizada = fecha_texto.replace("/", "-").replace(".", "-")
        for formato in ["%d-%m-%Y", "%d-%m-%y"]:
            try:
                fecha_obj = datetime.strptime(fecha_normalizada, formato).date()
                break
            except ValueError:
                continue

    # =========================================================================
    # 5. EXTRAER NOMBRE Y APELLIDOS (Estrategia Limpia de Primera Línea Válida)
    # =========================================================================
    nombres = "Candidato"
    apellido_paterno = "Desconocido"
    apellido_materno = None
    
    # Filtramos líneas que definitivamente son títulos del documento o ruidos
    ignorar_en_nombre = ["curriculum", "cv", "hoja de vida", "resumen", "perfil"]
    
    # Obtenemos las líneas del texto normalizado (donde ya reparamos a Cathy)
    lineas_limpias = []
    for l in texto_normalizado.split('\n'):
        l_strip = l.strip()
        if not l_strip:
            continue
        # Si la línea es solo datos de contacto, la saltamos
        if "@" in l_strip or "http" in l_strip.lower() or f"{telefono or 'XYZ'}" in l_strip.replace(" ", ""):
            continue
        # Si contiene títulos genéricos de CV, la saltamos
        if any(x in l_strip.lower() for x in ignorar_en_nombre) and len(l_strip.split()) < 4:
            continue
        lineas_limpias.append(l_strip)

    if lineas_limpias:
        # Asumimos con alta confianza que la primera línea limpia de verdad es el nombre
        nombre_completo = lineas_limpias[0]
        
        # Quitamos caracteres extraños como dos puntos o guiones que queden al final
        nombre_completo = re.sub(r'[:\-,_]', '', nombre_completo).strip()
        palabras = nombre_completo.split()
        
        if len(palabras) == 2:
            nombres = palabras[0]
            apellido_paterno = palabras[1]
        elif len(palabras) == 3:
            nombres = palabras[0]
            apellido_paterno = palabras[1]
            apellido_materno = palabras[2]
        elif len(palabras) >= 4:
            # Caso típico chileno: "Juan Pablo Pérez Gómez"
            nombres = f"{palabras[0]} {palabras[1]}"
            apellido_paterno = palabras[2]
            apellido_materno = palabras[3]

    # =========================================================================
    # 6. EXTRAER URLS (LINKEDIN / GITHUB)
    # =========================================================================
    linkedin_url = None
    github_url = None
    match_li = re.search(r'(https?://)?(www\.)?linkedin\.com/in/[a-zA-Z0-9_-]+', texto_normalizado, re.IGNORECASE)
    if match_li:
        linkedin_url = match_li.group(0)
    match_gh = re.search(r'(https?://)?(www\.)?github\.com/[a-zA-Z0-9_-]+', texto_normalizado, re.IGNORECASE)
    if match_gh:
        github_url = match_gh.group(0)

    # =========================================================================
    # 7. OPTIMIZACIÓN DEL RESUMEN PROFESIONAL
    # =========================================================================
    lineas_reales = []
    for l in texto_normalizado.split('\n'):
        l_limpia = l.strip()
        if len(l_limpia.split()) > 4 and "@" not in l_limpia and "http" not in l_limpia:
            if any(x in l_limpia.lower() for x in ["acerca de mi", "perfil profesional", "extracto"]):
                continue
            lineas_reales.append(l_limpia)
            
    if len(lineas_reales) >= 2:
        resumen_profesional = " ".join(lineas_reales[:3])
    else:
        resumen_lineas = lineas_limpias[1:4] if len(lineas_limpias) > 4 else lineas_limpias
        resumen_profesional = " ".join(resumen_lineas)

    resumen_profesional = re.sub(r'\s+', ' ', resumen_profesional).strip()
    resumen_profesional = resumen_profesional[:250] + "..."

    return CandidatoCreate(
        nombres=nombres,
        apellido_paterno=apellido_paterno,
        apellido_materno=apellido_materno,
        correo_electronico=correo,
        telefono_contacto=telefono,
        rut_candidato=rut,
        fecha_nacimiento=fecha_obj,
        linkedin_url=linkedin_url,
        github_url=github_url,
        pretension_renta=None,
        disponibilidad=None,
        resumen_profesional=resumen_profesional
    )
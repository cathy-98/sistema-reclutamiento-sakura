import io
import re
from pypdf import PdfReader
from fastapi import HTTPException
from app.schemas import CandidatoCreate

def extraer_texto_pdf(contenido_archivo: bytes) -> str:
    """
    Lee los bytes de un archivo PDF en memoria y extrae todo su texto plano utilizando pypdf.
    """
    try:
        pdf_file = io.BytesIO(contenido_archivo)
        reader = PdfReader(pdf_file)
        texto_completo = []
        
        for pagina in reader.pages:
            texto_pagina = pagina.extract_text()
            if texto_pagina:
                texto_completo.append(texto_pagina)
                
        texto = "\n".join(texto_completo)

        texto = limpiar_texto(texto)

        return texto
    except Exception as e:
        raise HTTPException(
            status_code=400,
            detail=f"No se pudo extraer el texto del archivo PDF: {str(e)}"
        )
    
def limpiar_texto(texto):
    """
    Elimina caracteres nulos (\x00) y espacios sobrantes que pueden
    provocar errores al guardar en PostgreSQL.
    """

    if texto is None:
        return None

    return (
        str(texto)
        .replace("\x00", "")
        .replace("\u0000", "")
        .replace("\r", "")
        .strip()
    )

def extraer_texto_docx(contenido_archivo: bytes) -> str:
    """
    Lee los bytes de un archivo DOCX en memoria y extrae todo su texto plano utilizando python-docx.
    """
    try:
        import docx
        docx_file = io.BytesIO(contenido_archivo)
        doc = docx.Document(docx_file)
        texto_completo = []
        
        for parrafo in doc.paragraphs:
            if parrafo.text:
                texto_completo.append(parrafo.text)
                
        for tabla in doc.tables:
            for fila in tabla.rows:
                for celda in fila.cells:
                    if celda.text:
                        texto_completo.append(celda.text)
                        
        texto = "\n".join(texto_completo)

        texto = limpiar_texto(texto)

        return texto
    except ImportError:
        raise HTTPException(
            status_code=502,
            detail="La librería 'python-docx' no está instalada. Por favor ejecute: pip install python-docx"
        )
    except Exception as e:
        raise HTTPException(
            status_code=400,
            detail=f"No se pudo extraer el texto del archivo DOCX: {str(e)}"
        )

def limpiar_texto(texto):
    """
    Elimina caracteres nulos (\x00) y espacios sobrantes que pueden
    provocar errores al guardar en PostgreSQL.
    """

    if texto is None:
        return None

    return (
        str(texto)
        .replace("\x00", "")
        .replace("\u0000", "")
        .replace("\r", "")
        .strip()
    )
    
def nombre_desde_correo(correo):

    if not correo:
        return None

    usuario = correo.split("@")[0]

    usuario = re.sub(r"[0-9]", "", usuario)

    partes = re.split(r"[._\-]", usuario)

    partes = [p.capitalize() for p in partes if p]

    if len(partes) >= 2:
        return partes

    return None


def extraer_nombre_completo(texto_cv, correo=None):

    lineas = [
        l.strip()
        for l in texto_cv.split("\n")
        if l.strip()
    ]

    palabras_prohibidas = [

        "curriculum",
        "currículum",
        "cv",
        "resume",
        "hoja",
        "vida",
        "perfil",
        "profesional",

        "ingeniero",
        "ingeniera",
        "analista",
        "developer",
        "desarrollador",
        "programador",
        "consultor",
        "software",

        "correo",
        "email",
        "contacto",
        "linkedin",
        "github",
        "telefono",
        "teléfono",

        "experiencia",
        "educación",
        "educacion",
        "habilidades",
        "competencias"
    ]

    mejor_linea = None
    mejor_puntaje = -1

    for indice, linea in enumerate(lineas[:15]):

        puntaje = 0

        texto = linea.lower()

        if any(x in texto for x in palabras_prohibidas):
            continue

        # Debe contener solamente letras y espacios
        if re.match(
            r"^[A-Za-zÁÉÍÓÚÜÑáéíóúüñ\s]+$",
            linea
        ):
            puntaje += 4
        else:
            continue

        palabras = linea.split()

        # Cantidad razonable de palabras para nombre completo
        if 3 <= len(palabras) <= 5:
            puntaje += 4
        else:
            continue

        # Mientras más arriba aparezca, más puntos
        puntaje += max(0, 5 - indice)

        # Muchas veces el nombre viene en mayúsculas
        if linea == linea.upper():
            puntaje += 2

        # Evitar líneas excesivamente largas
        if len(linea) < 40:
            puntaje += 1

        if puntaje > mejor_puntaje:
            mejor_puntaje = puntaje
            mejor_linea = linea

    # Si encontró una buena línea

    if mejor_linea:

        palabras = mejor_linea.title().split()

        apellido_materno = palabras[-1]
        apellido_paterno = palabras[-2]
        nombres = " ".join(palabras[:-2])

        return (
            nombres,
            apellido_paterno,
            apellido_materno
        )

    # Intentar reconstruir desde el correo

    posible = nombre_desde_correo(correo)

    if posible:

        if len(posible) == 2:

            return (
                posible[0],
                posible[1],
                None
            )

        if len(posible) >= 3:

            return (
                " ".join(posible[:-2]),
                posible[-2],
                posible[-1]
            )

    return (
        "Candidato",
        "Desconocido",
        None
    )

async def procesar_cv_con_python(texto_cv: str) -> CandidatoCreate:
    """
    Procesa el texto del CV utilizando únicamente expresiones regulares
    y lógica nativa de Python.
    """

    # =========================================================================
    # 1. EXTRAER CORREO ELECTRÓNICO
    # =========================================================================

    pattern_email = r'[a-zA-Z0-9\-\.]+@[a-zA-Z0-9\-\.]+\.[a-zA-Z]{2,5}'

    match_email = re.search(pattern_email, texto_cv)

    if match_email:
        correo = match_email.group(0).strip().lower()

    else:
        correo = None

        for linea in texto_cv.split("\n"):

            if "@" in linea:

                linea_sin_espacios = linea.replace(" ", "")

                match_linea = re.search(
                    pattern_email,
                    linea_sin_espacios
                )

                if match_linea:
                    correo = match_linea.group(0).strip().lower()
                    break

    if not correo:
        raise HTTPException(
            status_code=422,
            detail="No se pudo procesar el CV: No se encontró un correo electrónico válido."
        )

    # =========================================================================
    # 2. EXTRAER TELÉFONO
    # =========================================================================

    pattern_tel = r'\+?\d[\d\s-]{7,13}\d'

    matches_tel = re.findall(
        pattern_tel,
        texto_cv
    )

    telefono = None

    for tel in matches_tel:

        limpio = tel.replace(" ", "").replace("-", "")

        if "569" in limpio or len(limpio) >= 9:
            telefono = limpio
            break

    # =========================================================================
    # 3. EXTRAER NOMBRE Y APELLIDOS
    # =========================================================================

    lineas = [
        l.strip()
        for l in texto_cv.split("\n")
        if l.strip()
    ]

    nombres, apellido_paterno, apellido_materno = extraer_nombre_completo(
        texto_cv,
        correo
    )

    nombres = nombres.title()
    apellido_paterno = apellido_paterno.title()

    if apellido_materno:
        apellido_materno = apellido_materno.title()

    # =========================================================================
    # 4. EXTRAER LINKEDIN
    # =========================================================================

    linkedin_url = None

    match_li = re.search(
        r'(https?://)?(www\.)?linkedin\.com/in/[a-zA-Z0-9_-]+',
        texto_cv,
        re.IGNORECASE
    )

    if match_li:
        linkedin_url = match_li.group(0)

    else:

        match_li_short = re.search(
            r'linkedin:\s*/([a-zA-Z0-9_-]+)',
            texto_cv,
            re.IGNORECASE
        )

        if match_li_short:
            linkedin_url = (
                f"https://www.linkedin.com/in/"
                f"{match_li_short.group(1)}"
            )

    # =========================================================================
    # 5. EXTRAER GITHUB
    # =========================================================================

    github_url = None

    match_gh = re.search(
        r'(https?://)?(www\.)?github\.com/[a-zA-Z0-9_-]+',
        texto_cv,
        re.IGNORECASE
    )

    if match_gh:
        github_url = match_gh.group(0)

    # =========================================================================
    # 6. RESUMEN PROFESIONAL
    # =========================================================================

    inicio = 0

    for i, linea in enumerate(lineas):

        if nombres.lower() in linea.lower():
            inicio = i + 1
            break

    resumen_lineas = lineas[inicio:inicio + 4]

    resumen_profesional = " ".join(
        resumen_lineas
    )[:250]

    if resumen_profesional:
        resumen_profesional += "..."

    # =========================================================================
    # 7. RETORNAR OBJETO
    # =========================================================================
# ==========================================================
    # LIMPIAR CARACTERES NULOS PARA POSTGRESQL
    # ==========================================================

    nombres = limpiar_texto(nombres)

    apellido_paterno = limpiar_texto(
        apellido_paterno
    )

    apellido_materno = limpiar_texto(
        apellido_materno
    )

    correo = limpiar_texto(
        correo
    )

    telefono = limpiar_texto(
        telefono
    )

    linkedin_url = limpiar_texto(
        linkedin_url
    )

    github_url = limpiar_texto(
        github_url
    )

    resumen_profesional = limpiar_texto(
        resumen_profesional
    )

    print("===================================")
    print("NOMBRE:", repr(nombres))
    print("APELLIDO P:", repr(apellido_paterno))
    print("APELLIDO M:", repr(apellido_materno))
    print("EMAIL:", repr(correo))
    print("RESUMEN:", repr(resumen_profesional))
    print("===================================")

    return CandidatoCreate(

        nombres=nombres,

        apellido_paterno=apellido_paterno,

        apellido_materno=apellido_materno,

        correo_electronico=correo,

        telefono_contacto=telefono,

        rut_candidato=None,

        fecha_nacimiento=None,

        linkedin_url=linkedin_url,

        github_url=github_url,

        pretension_renta=None,

        disponibilidad=None,

        resumen_profesional=resumen_profesional

    )